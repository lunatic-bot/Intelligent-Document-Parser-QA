# tests/generate_sample_files.py
import os
from docx import Document
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from PIL import Image, ImageDraw, ImageFont

SAMPLE_DIR = "tests/sample_files"


def create_txt():
    with open(os.path.join(SAMPLE_DIR, "sample.txt"), "w", encoding="utf-8") as f:
        f.write("This is a simple text file for testing.\nIt contains multiple lines.\n")


def create_docx():
    doc = Document()
    doc.add_paragraph("This is a DOCX test file.")
    doc.add_paragraph("It contains two paragraphs.")
    doc.save(os.path.join(SAMPLE_DIR, "sample.docx"))


def create_pdf():
    c = canvas.Canvas(os.path.join(SAMPLE_DIR, "sample.pdf"), pagesize=letter)
    c.drawString(100, 750, "This is a test PDF.")
    c.drawString(100, 735, "It has two lines of text.")
    c.save()


def create_image_pdf():
    img = Image.new("RGB", (600, 100), color=(255, 255, 255))
    draw = ImageDraw.Draw(img)
    draw.text((10, 40), "This is an OCR image PDF.", fill=(0, 0, 0))
    image_path = os.path.join(SAMPLE_DIR, "ocr_sample.png")
    img.save(image_path)

    c = canvas.Canvas(os.path.join(SAMPLE_DIR, "sample_image.pdf"), pagesize=letter)
    c.drawImage(image_path, 100, 600, width=400, height=100)
    c.save()

    os.remove(image_path)


def ensure_sample_files():
    os.makedirs(SAMPLE_DIR, exist_ok=True)
    create_txt()
    create_docx()
    create_pdf()
    create_image_pdf()


if __name__ == "__main__":
    ensure_sample_files()
    print("Sample test documents generated including OCR PDF.")
